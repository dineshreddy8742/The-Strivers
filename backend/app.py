import os
import base64
import re
from io import BytesIO
from flask import Flask, request, jsonify
from flask_cors import CORS
from dotenv import load_dotenv
from groq import Groq
import google.generativeai as genai
from gtts import gTTS
import concurrent.futures
import time
import PyPDF2
import docx
from werkzeug.utils import secure_filename

# --- SETUP ---
load_dotenv()
app = Flask(__name__)
CORS(app)

# Configure file uploads
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'rtf'}

# Initialize API Clients
try:
    groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    print("✅ API clients initialized successfully")
except Exception as e:
    print(f"❌ Error initializing API clients: {e}")

# Enhanced LLM Models for different purposes
LLM_MODELS = {
    "basic": {
        "model": "llama-3.1-8b-instant",
        "max_tokens": 150,
        "temperature": 0.5,
        "description": "Quick and simple responses"
    },
    "research": {
        "model": "llama-3.3-70b-versatile", 
        "max_tokens": 800,
        "temperature": 0.3,
        "description": "Deep research and comprehensive analysis"
    },
    "creative": {
        "model": "mixtral-8x7b-32768",
        "max_tokens": 400,
        "temperature": 0.8,
        "description": "Creative content and diagrams"
    },
    "technical": {
        "model": "llama-3.3-70b-versatile",
        "max_tokens": 600,
        "temperature": 0.2,
        "description": "Technical explanations and code"
    }
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file):
    """Enhanced file text extraction with PDF and DOCX support"""
    try:
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        if file_extension == 'txt':
            try:
                content = file.read()
                for encoding in ['utf-8', 'latin-1', 'ascii', 'cp1252']:
                    try:
                        return content.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                return content.decode('utf-8', errors='ignore')
            except Exception as e:
                return f"Error reading text file: {str(e)}"
        
        elif file_extension == 'pdf':
            try:
                file.seek(0)
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as page_error:
                        text += f"\n[Error reading page {page_num + 1}: {page_error}]\n"
                return text.strip() if text.strip() else "No readable text found in PDF"
            except Exception as e:
                return f"Error reading PDF file: {str(e)}"
        
        elif file_extension in ['docx', 'doc']:
            try:
                file.seek(0)
                doc = docx.Document(file)
                text = ""
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n\n"
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text += " | ".join(row_text) + "\n"
                
                return text.strip() if text.strip() else "No readable text found in document"
            except Exception as e:
                return f"Error reading Word document: {str(e)}"
        
        else:
            available_formats = list(ALLOWED_EXTENSIONS)
            return f"Unsupported file format. Supported formats: {', '.join(available_formats)}"
        
    except Exception as e:
        return f"Error processing file: {str(e)}"

def select_model_and_prompt(query, mode="basic", content_type="general"):
    """Enhanced model selection based on mode and content type"""
    
    if mode == "research":
        if any(keyword in query.lower() for keyword in ['flow', 'diagram', 'flowchart', 'process', 'workflow']):
            return LLM_MODELS["creative"], get_diagram_prompt(query)
        else:
            return LLM_MODELS["research"], get_research_prompt(query)
    
    elif mode == "basic":
        return LLM_MODELS["basic"], get_basic_prompt(query)
    
    else:  # Auto-detect mode
        if any(keyword in query.lower() for keyword in ['explain in detail', 'comprehensive', 'research', 'analyze thoroughly']):
            return LLM_MODELS["research"], get_research_prompt(query)
        elif any(keyword in query.lower() for keyword in ['flow', 'diagram', 'flowchart']):
            return LLM_MODELS["creative"], get_diagram_prompt(query)
        else:
            return LLM_MODELS["basic"], get_basic_prompt(query)

def get_basic_prompt(query):
    return f"""You are a helpful sustainability assistant. Provide a clear, concise answer to this question about environmental science, metals, mining, or sustainability.

Keep your response:
- Under 100 words
- Easy to understand
- Practical and actionable
- Well-structured with bullet points if needed

Question: {query}"""

def get_research_prompt(query):
    return f"""You are an expert sustainability researcher with deep knowledge in environmental science, metals, mining, and sustainable practices.

Provide a comprehensive, well-researched response that includes:
- Detailed explanation of the topic
- Current industry standards and best practices
- Environmental impact considerations  
- Latest developments and innovations
- Practical recommendations
- Supporting data or statistics where relevant

Structure your response with clear headings and subpoints. Aim for 400-600 words.

Research Query: {query}"""

def get_diagram_prompt(query):
    return f"""You are a process flow expert. Create a detailed textual representation of the requested flow diagram or process.

Format your response as:
1. Brief overview of the process
2. Step-by-step flow with clear connections
3. Key decision points and branches
4. Important notes or considerations

Use ASCII characters, arrows (→, ↓, ↑), and clear formatting to create a visual flow.

Flow Diagram Request: {query}"""

def detect_content_type(query):
    """Detect what type of content the user is requesting"""
    query_lower = query.lower()
    
    if any(keyword in query_lower for keyword in ['image', 'picture', 'photo', 'visual', 'generate', 'create', 'draw']):
        return "image"
    elif any(keyword in query_lower for keyword in ['flow', 'diagram', 'flowchart', 'process map', 'workflow']):
        return "diagram" 
    elif any(keyword in query_lower for keyword in ['research', 'analyze', 'detailed', 'comprehensive', 'in-depth']):
        return "research"
    else:
        return "general"

def clean_response_formatting(text):
    """Enhanced response formatting"""
    if not text:
        return text
    
    # Remove asterisks and improve formatting
    text = re.sub(r'\*{2,}', '', text)
    text = re.sub(r'\*([^*]+)\*', r'<strong>\1</strong>', text)
    text = text.replace('*', '')
    
    # Enhanced list formatting
    lines = text.split('\n')
    formatted_lines = []
    in_list = False
    
    for line in lines:
        line = line.strip()
        
        # Handle headers (## or ###)
        if line.startswith('##'):
            if in_list:
                formatted_lines.append('</ul>')
                in_list = False
            formatted_lines.append(f'<h3>{line.replace("#", "").strip()}</h3>')
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('• ') or re.match(r'^\d+\.', line):
            if not in_list:
                formatted_lines.append('<ul>')
                in_list = True
            list_text = re.sub(r'^[\-•\d\.]\s*', '', line)
            formatted_lines.append(f'<li>{list_text}</li>')
        
        # Handle regular paragraphs
        else:
            if in_list:
                formatted_lines.append('</ul>')
                in_list = False
            if line:
                formatted_lines.append(f'<p>{line}</p>')
    
    if in_list:
        formatted_lines.append('</ul>')
    
    result = '\n'.join(formatted_lines)
    return result if result.strip() else f'<p>{text}</p>'

def generate_tts_audio(text, lang='en'):
    """Generate TTS audio with error handling"""
    try:
        clean_text = re.sub(r'<[^>]+>', '', text)
        if not clean_text.strip() or len(clean_text) > 1000:  # Skip very long texts
            return None
            
        tts = gTTS(text=clean_text[:800], lang=lang, slow=False)  # Limit length
        mp3_fp = BytesIO()
        tts.write_to_fp(mp3_fp)
        mp3_fp.seek(0)
        return base64.b64encode(mp3_fp.read()).decode('utf-8')
    except Exception as e:
        print(f"TTS Error: {e}")
        return None

# --- API ENDPOINTS ---
@app.route('/')
def index():
    return "🚀 Advanced Eco AI Assistant - Multi-Modal LLM System"

@app.route('/health')
def health_check():
    return jsonify({
        "status": "healthy",
        "models": list(LLM_MODELS.keys()),
        "supported_formats": list(ALLOWED_EXTENSIONS),
        "max_file_size": "50MB",
        "features": ["basic_mode", "research_mode", "image_generation", "flow_diagrams", "pdf_support"],
        "timestamp": time.time()
    })

@app.route('/chat', methods=['POST'])
def chat_handler():
    data = request.json
    user_text = data.get('message', '')
    history = data.get('history', [])
    mode = data.get('mode', 'basic')  # 'basic' or 'research' 
    
    if not user_text:
        return jsonify({"error": "No message provided"}), 400

    try:
        # Detect content type and select appropriate model
        content_type = detect_content_type(user_text)
        model_config, system_prompt = select_model_and_prompt(user_text, mode, content_type)
        
        print(f"🤖 Mode: {mode}, Content Type: {content_type}, Model: {model_config['model']}")
        
        # Build conversation context
        messages = [{"role": "system", "content": system_prompt}]
        
        # Add recent history (limit based on mode)
        history_limit = 3 if mode == "basic" else 6
        recent_history = history[-history_limit:] if len(history) > history_limit else history
        
        for entry in recent_history:
            if isinstance(entry, dict) and 'user' in entry and 'ai' in entry:
                messages.append({"role": "user", "content": str(entry['user'])})
                messages.append({"role": "assistant", "content": str(entry['ai'])})
        
        messages.append({"role": "user", "content": user_text})

        # Get AI response with selected model
        chat_completion = groq_client.chat.completions.create(
            messages=messages,
            model=model_config["model"],
            max_tokens=model_config["max_tokens"],
            temperature=model_config["temperature"]
        )
        
        raw_reply = chat_completion.choices[0].message.content
        clean_reply = clean_response_formatting(raw_reply)
        
        # Generate TTS only for shorter responses
        audio_base64 = None
        if mode == "basic" and len(raw_reply) < 500:
            with concurrent.futures.ThreadPoolExecutor() as executor:
                tts_future = executor.submit(generate_tts_audio, raw_reply)
                try:
                    audio_base64 = tts_future.result(timeout=5)
                except:
                    pass

        # Check for image generation request
        image_prompt = None
        if content_type == "image":
            # Extract image subject
            clean_prompt = re.sub(r'(generate|create|make|show)\s+(image|picture|photo)\s+(of|showing|depicting)\s*', '', user_text.lower())
            image_prompt = f"Professional high-quality illustration: {clean_prompt}"

        return jsonify({
            "reply_text": clean_reply,
            "reply_audio": audio_base64,
            "image_prompt": image_prompt,
            "mode_used": mode,
            "model_used": model_config["model"],
            "content_type": content_type,
            "auto_generate_image": content_type == "image"
        })
        
    except Exception as e:
        print(f"Chat handler error: {e}")
        return jsonify({"error": f"Chat processing failed: {str(e)}"}), 500

@app.route('/upload-file', methods=['POST'])
def upload_file_handler():
    """Enhanced file upload with better PDF/DOCX support"""
    print("📁 File upload endpoint accessed")
    
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    if not allowed_file(file.filename):
        supported = ', '.join(ALLOWED_EXTENSIONS)
        return jsonify({"error": f"File type not supported. Supported: {supported}"}), 400
    
    try:
        print(f"📄 Processing file: {file.filename}")
        file_content = extract_text_from_file(file)
        
        if not file_content or file_content.startswith("Error"):
            return jsonify({"error": file_content or "Could not extract text from file"}), 400
        
        # Enhanced analysis using research model
        analysis_prompt = f"""Analyze this document comprehensively and provide:

1. **Document Summary** (2-3 sentences)
2. **Key Topics** (bullet points)
3. **Sustainability & Environmental Aspects** (if any)
4. **Important Data/Metrics** (if any)
5. **Actionable Insights** (recommendations)

Document Content:
{file_content[:4000]}"""
        
        messages = [
            {"role": "system", "content": "You are an expert document analyzer specializing in sustainability and environmental science. Provide structured analysis with clear headings."},
            {"role": "user", "content": analysis_prompt}
        ]
        
        response = groq_client.chat.completions.create(
            messages=messages,
            model=LLM_MODELS["research"]["model"],
            max_tokens=600,
            temperature=0.3
        )
        
        analysis = clean_response_formatting(response.choices[0].message.content)
        print("✅ Document analysis completed")
        
        return jsonify({
            "filename": file.filename,
            "file_content": file_content[:1500] + "..." if len(file_content) > 1500 else file_content,
            "analysis": analysis,
            "content_length": len(file_content),
            "file_type": file.filename.split('.')[-1].upper(),
            "status": "success"
        })
        
    except Exception as e:
        print(f"❌ File processing error: {e}")
        return jsonify({"error": f"File processing failed: {str(e)}"}), 500

@app.route('/transcribe', methods=['POST'])
def transcribe_handler():
    if 'audio_data' not in request.files:
        return jsonify({"error": "No audio file provided"}), 400
    
    audio_file = request.files['audio_data']
    
    try:
        audio_file.seek(0)
        
        transcription = groq_client.audio.transcriptions.create(
            file=("recording.webm", audio_file.read(), "audio/webm"),
            model="whisper-large-v3-turbo",
            language="en",
            response_format="text",
            temperature=0.0
        )
        
        clean_transcription = transcription.strip() if transcription else ""
        
        if not clean_transcription:
            return jsonify({"error": "No speech detected in audio"}), 400
            
        return jsonify({
            "transcribed_text": clean_transcription,
            "show_transcription": False  # Hide by default
        })
        
    except Exception as e:
        print(f"Transcription error: {e}")
        return jsonify({"error": f"Transcription failed: {str(e)}"}), 500

@app.route('/generate-image', methods=['POST'])
def generate_image_handler():
    prompt = request.json.get('prompt', '')
    style = request.json.get('style', 'professional')
    
    if not prompt:
        return jsonify({"error": "No image prompt provided"}), 400

    try:
        # Enhanced image generation with multiple sources
        keywords = extract_image_keywords(prompt)
        
        # Try different image services
        image_sources = [
            f"https://source.unsplash.com/1024x768/?{','.join(keywords)}",
            f"https://picsum.photos/1024/768?random={int(time.time())}",
            "https://via.placeholder.com/1024x768/667eea/FFFFFF?text=Generated+Visualization"
        ]
        
        for url in image_sources:
            try:
                import requests
                response = requests.get(url, timeout=10)
                if response.status_code == 200:
                    return jsonify({
                        "image_url": url, 
                        "keywords": keywords,
                        "style": style,
                        "source": "unsplash" if "unsplash" in url else "fallback"
                    })
            except:
                continue
        
        return jsonify({"error": "Image generation temporarily unavailable"}), 500
        
    except Exception as e:
        return jsonify({"error": f"Image generation failed: {str(e)}"}), 500

def extract_image_keywords(prompt):
    """Enhanced keyword extraction for better images"""
    # Remove common phrases
    clean_prompt = re.sub(r'(professional|high-quality|illustration|image|picture|photo|of|showing|depicting)', '', prompt.lower())
    
    # Extract meaningful words
    words = re.findall(r'\w+', clean_prompt)
    keywords = [word for word in words if len(word) > 3]
    
    # Prioritize sustainability and tech terms
    priority_terms = ['sustainability', 'environment', 'renewable', 'solar', 'wind', 'mining', 'technology', 'green', 'eco']
    priority_keywords = [word for word in keywords if word in priority_terms][:2]
    
    # Add other relevant terms
    if len(priority_keywords) < 3:
        priority_keywords.extend([word for word in keywords if word not in priority_keywords][:3-len(priority_keywords)])
    
    return priority_keywords or ['sustainability', 'technology', 'environment']

if __name__ == "__main__":
    print("🚀 Starting Advanced Eco AI Assistant...")
    print(f"🧠 Available Models: {list(LLM_MODELS.keys())}")
    print(f"📄 Supported Formats: {', '.join(ALLOWED_EXTENSIONS)}")
    print("⚡ Features: Research Mode, Image Generation, PDF Support, Flow Diagrams")
    app.run(host='127.0.0.1', port=5000, debug=True, threaded=True)
